from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.shared import Pt


def add_heading(document: Document, text: str) -> None:
    document.add_paragraph(text, style="Heading 1")


def add_bullet_list(document: Document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_labeled_bullets(document: Document, labeled_items):
    for label, description in labeled_items:
        paragraph = document.add_paragraph(style="List Bullet")
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(description)


def add_numbered_list(document: Document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_code_block(document: Document, code: str):
    paragraph = document.add_paragraph()
    for line in code.splitlines():
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def main():
    report_path = Path("food_tracker_report.docx")

    document = Document()
    document.add_heading("Отчёт по проекту «Food Tracker»", level=1)

    # описание цели проекта
    add_heading(document, "1. Постановка задачи")
    document.add_paragraph(
        "Цель проекта — автоматизировать сбор и сравнение цен на продукты повседневного "
        "спроса в разных розничных сетях. Приложение объединяет данные, полученные из "
        "интернет-витрин «Окей» и «Светофор», нормализует категории, ищет одинаковые и "
        "похожие позиции и показывает покупателю, где нужный товар стоит дешевле."
    )
    document.add_paragraph(
        "Пользователь получает web-интерфейс с подборкой товаров дня, навигацией по "
        "категориям и полнотекстовым поиском. За кулисами работают парсеры, база данных "
        "и алгоритмы сравнения на основе нечёткого поиска."
    )

    # список технологий стека
    add_heading(document, "2. Использованные технологии")
    technologies = [
        "Python 3.x для серверной логики и скриптов парсинга",
        "Flask с blueprint-архитектурой для веб-приложения",
        "Flask-SQLAlchemy + SQLite для ORM-уровня и постоянного хранилища",
        "Jinja2, HTML5 и CSS3 для пользовательского интерфейса",
        "Selenium + WebDriver Manager для сбора данных с сайтов ретейлеров",
        "RapidFuzz и FuzzyWuzzy для поиска похожих товаров",
        "Git и виртуальное окружение venv для управления зависимостями и версий",
    ]
    add_bullet_list(document, technologies)

    # описание библиотек с их назначением
    add_heading(document, "3. Библиотеки и их назначение")
    libraries = [
        ("Flask", "поднимает веб-сервер, обрабатывает маршруты и рендерит шаблоны"),
        ("Flask-SQLAlchemy", "даёт объектно-реляционное отображение и управление сессиями БД"),
        (
            "SQLite",
            "хранит нормализованные магазины, категории и товары внутри файла data/food_tracker.db",
        ),
        ("Jinja2", "шаблонизатор для страниц `frontend/templates`"),
        ("Selenium", "парсит карточки товаров на страницах магазинов"),
        ("webdriver-manager", "автоматически загружает совместимые ChromeDriver'ы для Selenium"),
        ("BeautifulSoup4 (bs4)", "парсит HTML-разметку страниц магазина «Светофор» для извлечения данных о товарах"),
        ("RapidFuzz", "быстро считает метрики похожести строк при объединении товаров"),
        ("FuzzyWuzzy + python-Levenshtein", "расширяют пул алгоритмов нечёткого поиска"),
        ("SQLAlchemy.func", "используется в поиске для case-insensitive запросов по названию товара"),
        ("python-docx", "сгенерировал этот отчёт в формате .docx"),
    ]
    add_labeled_bullets(document, libraries)

    # хронология разработки
    add_heading(document, "4. Этапы разработки проекта")
    steps = [
        "Исследование предметной области, выделение ключевых категорий товаров и источников данных.",
        "Реализация парсеров для витрин «Окей» и «Светофор», выгрузка позиций в отдельные SQLite-базы.",
        "Нормализация категорий и объединение разных источников в общую БД через скрипт `merge_to_main_db.py`.",
        "Проектирование моделей `Store`, `Category`, `Product` и настройка ORM-слоя.",
        "Создание Flask-приложения с блупринтами `main` и `search_`, конфигурация шаблонов и статики.",
        "Добавление логики по поиску похожих товаров и подсветке позиций, присутствующих в обоих магазинах.",
        "Разработка пользовательского интерфейса (главная, категории, карточки, поиск, страница «О нас»).",
        "Тестирование, наполнение базы и подготовка окружения к демонстрации.",
    ]
    add_numbered_list(document, steps)

    # пример кода с пояснениями
    add_heading(document, "5. Код с комментариями")
    document.add_paragraph(
        "Фрагмент обработчика категории иллюстрирует этапы группировки товаров, поиск "
        "совпадений между магазинами и формирование выдачи для шаблона."
    )
    code_snippet = dedent(
        """
        @main.route("/category/<int:category_id>")
        def category(category_id):
            # Получаем саму категорию и весь её ассортимент
            category_obj = Category.query.get_or_404(category_id)
            products = Product.query.filter_by(category_id=category_id).all()

            # Группируем товары по нормализованному названию
            by_norm = {}
            for product in products:
                key = _normalize_product_name(product.name)
                by_norm.setdefault(key, []).append(product)

            in_both = []
            remaining = []
            for normalized_name, grouped in by_norm.items():
                stores = {p.store.name for p in grouped}
                if len(stores) >= 2:
                    in_both.append({"name": grouped[0].name, "products": grouped})  # товар есть в двух сетях
                else:
                    remaining.extend(grouped)

            threshold_min = 60  # нижний порог похожести
            threshold_max = 95  # всё, что выше, считаем тем же товаром
            remaining_sorted = sorted(remaining, key=lambda p: _normalize_product_name(p.name))
            used_ids = set()
            similar_groups = []

            for i, base_product in enumerate(remaining_sorted):
                if base_product.id in used_ids:
                    continue  # пропускаем уже сгруппированные товары

                group = [base_product]
                for candidate in remaining_sorted[i + 1:]:
                    if candidate.id in used_ids:
                        continue

                    score = _calculate_similarity(base_product.name, candidate.name)
                    if threshold_min <= score < threshold_max:
                        group.append(candidate)
                        used_ids.add(candidate.id)

                if len(group) > 1:
                    similar_groups.append({
                        "name": min(group, key=lambda p: len(p.name)).name,
                        "products": group,
                    })  # показываем группу похожих товаров

            unique_products = [p for p in remaining if p.id not in used_ids]
            return render_template(
                "category.html",
                category=category_obj,
                in_both=in_both,
                similar_groups=similar_groups,
                unique_products=unique_products,
            )
        """
    ).strip()
    add_code_block(document, code_snippet)

    document.save(report_path)
    print(f"Report generated at: {report_path.resolve()}")


if __name__ == "__main__":
    main()


